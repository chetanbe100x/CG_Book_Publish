import sys
import re
import copy
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

def iter_block_members(parent):
    parent_elm = parent.element.body if (hasattr(parent, "element") and hasattr(parent.element, "body")) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

def copy_table_element(src_table, dest_doc):
    tbl_clone = copy.deepcopy(src_table._tbl)
    
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    # Apply fonts and cantSplit properties
    for r in tbl_clone.findall(".//w:r", namespaces=ns):
        text_nodes = r.findall("w:t", namespaces=ns)
        if text_nodes:
            text_val = "".join([t.text for t in text_nodes if t.text])
            is_hindi = any(ord(char) > 127 for char in text_val)
            
            rPr = r.get_or_add_rPr()
            
            sz = rPr.find("w:sz", namespaces=ns)
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), "23" if is_hindi else "21")
            
            szCs = rPr.find("w:szCs", namespaces=ns)
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rPr.append(szCs)
            szCs.set(qn("w:val"), "23" if is_hindi else "21")
            
            rFonts = rPr.find("w:rFonts", namespaces=ns)
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            if is_hindi:
                rFonts.set(qn("w:hint"), "default")
                rFonts.set(qn("w:ascii"), "Mangal")
                rFonts.set(qn("w:hAnsi"), "Mangal")
                rFonts.set(qn("w:cs"), "Mangal")
            else:
                rFonts.set(qn("w:hint"), "default")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rFonts.set(qn("w:cs"), "Times New Roman")
                
    for tr in tbl_clone.findall(".//w:tr", namespaces=ns):
        trPr = tr.get_or_add_trPr()
        if trPr.find("w:cantSplit", namespaces=ns) is None:
            trPr.append(OxmlElement("w:cantSplit"))
            
    dest_doc.element.body.append(tbl_clone)


def set_run_font_complex(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

def add_mixed_run(paragraph, text, is_bold=False, is_italic=False, size_hindi=12.0, size_latin=11.0):
    tokens = []
    current_token = ""
    current_type = None
    
    for char in text:
        is_hindi_char = ('\u0900' <= char <= '\u097f')
        char_type = "hindi" if is_hindi_char else "latin"
        
        if char in " \t\r\n,.:;?!()₹%-0123456789/\\":
            if current_type is None:
                current_type = char_type
            char_type = current_type
            
        if char_type != current_type:
            if current_token:
                tokens.append((current_token, current_type))
            current_token = char
            current_type = char_type
        else:
            current_token += char
            
    if current_token:
        tokens.append((current_token, current_type))
        
    for token_text, token_type in tokens:
        run = paragraph.add_run(token_text)
        run.bold = is_bold
        run.italic = is_italic
        if token_type == "hindi":
            set_run_font_complex(run, "Mangal")
            run.font.size = Pt(size_hindi)
        else:
            set_run_font_complex(run, "Times New Roman")
            run.font.size = Pt(size_latin)

def create_options_table(doc, options):
    # Create 2x2 table for A, B, C, D
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Remove borders
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        
    # Prevent row splitting and set widths
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = trPr.find(qn("w:cantSplit"))
        if cantSplit is None:
            trPr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.width = Inches(3.0)
            
    # Fill cells
    cells = [table.cell(0, 0), table.cell(0, 1), table.cell(1, 0), table.cell(1, 1)]
    for idx, text in enumerate(options):
        p = cells[idx].paragraphs[0]
        p.style = "Normal"
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True # Keep table options with the dotted lines
        
        # Style option label
        match = re.match(r"^\(?([a-dE-He-hA-Da-d])\)?\s*(.*)", text)
        if match:
            label = match.group(1)
            val = match.group(2)
            run_lbl = p.add_run(f"({label}) ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = RGBColor(0x16, 0x88, 0x73) # Accent green
            set_run_font_complex(run_lbl, "Times New Roman")
            run_lbl.font.size = Pt(11.0)
            add_mixed_run(p, val, size_hindi=11.5, size_latin=10.5)
        else:
            add_mixed_run(p, text, size_hindi=11.5, size_latin=10.5)

def add_dotted_line(doc, index=0, is_last=False, keep_with_next=False, spacing_points=30.5):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(0)
    # The default pitch is 30.5 pt. A default text run height is ~11 pt.
    # Therefore, to achieve 30.5 pt spacing between consecutive lines,
    # we set space_after to Pt(spacing_points - 11.0) = 19.5 pt.
    if is_last:
        p.paragraph_format.space_after = Pt(20.0) # Visual gap after the question block
    else:
        p.paragraph_format.space_after = Pt(max(1.0, spacing_points - 11.0))
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep_with_next
    
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "4")
    # Alternate space to prevent MS Word border merging
    space_val = "1" if index % 2 == 0 else "2"
    bottom.set(qn("w:space"), space_val)
    bottom.set(qn("w:color"), "B7B7B7")

    between = borders.find(qn("w:between"))
    if between is None:
        between = OxmlElement("w:between")
        borders.append(between)
    between.set(qn("w:val"), "dotted")
    between.set(qn("w:sz"), "4")
    between.set(qn("w:space"), space_val)
    between.set(qn("w:color"), "B7B7B7")



def main():
    source_path = Path(r"d:\Experiments_Projects\CG_Book_Publish\Books\Class 12\ACCOUNTANCY\Input\class_12_accountancy_paginated.docx")
    dest_path = Path(r"d:\Experiments_Projects\CG_Book_Publish\Books\Class 12\ACCOUNTANCY\Input\class_12_accountancy_reconstructed.docx")
    
    if not source_path.is_file():
        print(f"Source file not found: {source_path}")
        sys.exit(1)
        
    doc = Document(source_path)
    out_doc = Document()
    
    # Setup page margins
    for section in out_doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    line_count = 0
    is_mcq = False
    
    # Iterate through elements (paragraphs & tables) sequentially
    elements = list(iter_block_members(doc))
    
    i = 0
    while i < len(elements):
        member = elements[i]
        
        # If Table, copy element directly (preserves formatting and properties)
        if isinstance(member, Table):
            copy_table_element(member, out_doc)
            i += 1
            continue
            
        p = member
        text = p.text.replace("\u200b", "").replace("\ufeff", "").strip()
        if not text:
            i += 1
            continue
            
        # Check if heading
        is_heading = False
        if re.search(r"^(अध्याय|Chapter|CHAPTER|भाग|Part|Part-)", text):
            is_heading = True
        elif len(text) < 120 and any(kw in text for kw in ["Marks", "numerical", "Objective", "Theory", "Questions", "Question"]):
            is_heading = True
            
        if is_heading:
            # Try to match line count directives
            line_match = re.search(r"प्रत्येक\s+प्रश्न\s*(पर|में)\s*(\d+)\s*लाइन", text)
            if line_match:
                line_count = int(line_match.group(2))
                is_mcq = False
            else:
                line_match = re.search(r"(\d+)\s*(लाइन|line)", text)
                if line_match:
                    line_count = int(line_match.group(1))
                    is_mcq = False
                else:
                    # Check if it specifies MCQ
                    if any(kw in text.casefold() for kw in ["वस्तुनिष्ठ", "mcq", "multiple choice", "वैकल्पिक", "objective"]):
                        line_count = 0
                        is_mcq = True
                    else:
                        # Apply defaults for specific types if line count not explicitly written
                        if any(kw in text.casefold() for kw in ["numerical", "practical", "क्रियात्मक", "न्यूमेरिकल"]):
                            line_count = 10
                        elif any(kw in text.casefold() for kw in ["theory", "descriptive", "वर्णनात्मक", "सिद्धांतिक"]):
                            line_count = 8
                        else:
                            line_count = 0
                        is_mcq = False
                    
            print(f"Header: '{text}' -> Classified: MCQ={is_mcq}, Dotted Lines={line_count}".encode("ascii", errors="backslashreplace").decode("ascii"))
            
            # Write heading
            h_p = out_doc.add_paragraph()
            is_chapter = "अध्याय" in text or "Chapter" in text
            if is_chapter:
                h_p.alignment = 1 # Center
                h_p.paragraph_format.page_break_before = True # Always start new chapter on a new page!
                add_mixed_run(h_p, text, is_bold=True, size_hindi=15.0, size_latin=14.0)
            else:
                h_p.alignment = 0 # Left
                add_mixed_run(h_p, text, is_bold=True, size_hindi=13.0, size_latin=12.0)
                
            i += 1
            continue
            
        # Check if question
        is_q = re.match(r"^(Q\d+\.|प्रश्न\s*:\s*\d*|प्रश्न\s+\d+|Q\d+)", text)
        if is_q:
            q_label = is_q.group(1)
            q_text = text[len(q_label):].strip()
            
            # Gather English translation if immediately following
            eng_text = ""
            if i + 1 < len(elements) and isinstance(elements[i+1], Paragraph):
                next_p = elements[i+1]
                next_text = next_p.text.replace("\u200b", "").replace("\ufeff", "").strip()
                # If next line is not a question and not an option, treat as English translation
                if next_text and not re.match(r"^(Q\d+\.|प्रश्न|[\u200b\s]*\([a-dE-He-hA-Da-d]\))", next_text):
                    eng_text = next_text
                    i += 1 # Consume English paragraph
                    
            # Write Hindi Question
            q_p = out_doc.add_paragraph()
            q_p.paragraph_format.keep_together = True
            q_p.paragraph_format.keep_with_next = True
            
            run_lbl = q_p.add_run(f"{q_label} ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = RGBColor(0x16, 0x88, 0x73) # Accent green
            set_run_font_complex(run_lbl, "Times New Roman")
            run_lbl.font.size = Pt(11.5)
            add_mixed_run(q_p, q_text, size_hindi=11.5, size_latin=11.0)
            
            # Write English Translation
            if eng_text:
                eng_p = out_doc.add_paragraph()
                eng_p.paragraph_format.left_indent = Pt(12) # Indent translation
                eng_p.paragraph_format.keep_together = True
                eng_p.paragraph_format.keep_with_next = True
                add_mixed_run(eng_p, eng_text, size_hindi=11.0, size_latin=10.0)
                for run in eng_p.runs:
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate gray
                
            # If MCQ section or options are immediately following, gather options
            next_cand_text = ""
            if i + 1 < len(elements) and isinstance(elements[i+1], Paragraph):
                next_cand_text = elements[i+1].text.replace("\u200b", "").strip()
                
            if is_mcq or re.match(r"^\([a-dE-He-hA-Da-d]\)", next_cand_text):
                options = []
                opt_count = 0
                while opt_count < 4 and i + 1 < len(elements):
                    cand_el = elements[i+1]
                    if not isinstance(cand_el, Paragraph):
                        break
                    cand_text = cand_el.text.replace("\u200b", "").replace("\ufeff", "").strip()
                    if cand_text and re.match(r"^\([a-dE-He-hA-Da-d]\)", cand_text):
                        options.append(cand_text)
                        opt_count += 1
                        i += 1
                    elif not cand_text:
                        i += 1 # skip empty lines
                    else:
                        break
                        
                if len(options) == 4:
                    create_options_table(out_doc, options)
                else:
                    # Write whatever options we gathered
                    for opt in options:
                        opt_p = out_doc.add_paragraph()
                        opt_p.paragraph_format.keep_together = True
                        opt_p.paragraph_format.keep_with_next = True
                        opt_p.add_run(opt)
                # MCQ: Append exactly 2 dotted lines after options grid with gap
                for line_idx in range(2):
                    is_last_line = (line_idx == 1)
                    keep = not is_last_line
                    add_dotted_line(out_doc, line_idx, is_last=is_last_line, keep_with_next=keep)
            else:
                # Non-MCQ question: append dotted answer lines
                if line_count > 0:
                    for line_idx in range(line_count):
                        is_last_line = (line_idx == line_count - 1)
                        if line_count <= 4:
                            keep = not is_last_line
                        else:
                            keep = (line_idx < 2)
                        add_dotted_line(out_doc, line_idx, is_last=is_last_line, keep_with_next=keep)
                        
            i += 1
            continue
            
        # Default copy for regular body text/instructions
        new_p = out_doc.add_paragraph()
        new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
        new_p.alignment = p.alignment
        add_mixed_run(new_p, text, size_hindi=12.0, size_latin=11.0)
                
        i += 1
        
    out_doc.save(dest_path)
    print(f"Reconstructed document successfully written to: {dest_path}")

if __name__ == "__main__":
    main()
