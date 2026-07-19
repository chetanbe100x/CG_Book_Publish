from __future__ import annotations

import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, name):
    p_elm = paragraph._p
    bm_id = name.split('_')[1]
    
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bm_id)
    bm_start.set(qn('w:name'), name)
    
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bm_id)
    
    p_elm.insert(0, bm_start)
    p_elm.append(bm_end)

def generate_toc(doc: Document, config: dict) -> None:
    title_text = config.get("title", "विषय सूची (Table of Contents)")
    title_font = config.get("title_font", "Mangal")
    title_size_pt = config.get("title_size_pt", 16)
    link_font = config.get("link_font", "Mangal")
    link_size_pt = config.get("link_size_pt", 12)
    link_color = config.get("link_color", "0563C1")
    
    cover_end_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        text_norm = p.text.strip().replace("–", "-").replace("—", "-")
        if text_norm.startswith("अध्याय 1") or text_norm.lower().startswith("chapter 1"):
            cover_end_idx = idx
            break
            
    if cover_end_idx == -1:
        print("[toc_generator] Could not find Chapter 1 start to insert Table of Contents.")
        return
        
    chapters = []
    bookmark_counter = 1
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if re.search(r"^(अध्याय|Chapter|CHAPTER|भाग-ब|Part B|Part-B)", text) and len(text) < 250:
            if any(kw in text.lower() for kw in ["questions", "प्रश्न", "marks", "अंक"]):
                continue
            hyperlinks = p._p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink")
            if hyperlinks:
                continue
            bm_name = f"ch_{bookmark_counter}"
            add_bookmark(p, bm_name)
            chapters.append((text, bm_name))
            bookmark_counter += 1
            
    print(f"[toc_generator] Bookmarked {len(chapters)} chapters.")
    
    insert_before_el = doc.paragraphs[cover_end_idx]._p
    target_body = doc.element.body
    
    def insert_para_xml(xml_el):
        pos = target_body.index(insert_before_el)
        target_body.insert(pos, xml_el)

    pb_para = OxmlElement('w:p')
    run = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run.append(br)
    pb_para.append(run)
    insert_para_xml(pb_para)
    
    title_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    kwn = OxmlElement('w:keepWithNext')
    pPr.append(kwn)
    title_para.append(pPr)
    
    run_t = OxmlElement('w:r')
    rPr_t = OxmlElement('w:rPr')
    rFonts_t = OxmlElement('w:rFonts')
    rFonts_t.set(qn('w:ascii'), title_font)
    rFonts_t.set(qn('w:hAnsi'), title_font)
    rFonts_t.set(qn('w:cs'), title_font)
    rPr_t.append(rFonts_t)
    
    sz_t = OxmlElement('w:sz')
    sz_t.set(qn('w:val'), str(int(title_size_pt * 2)))
    rPr_t.append(sz_t)
    szCs_t = OxmlElement('w:szCs')
    szCs_t.set(qn('w:val'), str(int(title_size_pt * 2)))
    rPr_t.append(szCs_t)
    
    bold_t = OxmlElement('w:b')
    rPr_t.append(bold_t)
    run_t.append(rPr_t)
    t_text = OxmlElement('w:t')
    t_text.text = title_text
    run_t.append(t_text)
    title_para.append(run_t)
    insert_para_xml(title_para)
    
    space_p = OxmlElement('w:p')
    insert_para_xml(space_p)
    
    for text, bm_name in chapters:
        link_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '120')
        spacing.set(qn('w:after'), '120')
        pPr.append(spacing)
        
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        pPr.append(ind)
        
        link_p.append(pPr)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bm_name)
        hyperlink.set(qn('w:history'), '1')
        
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), link_color)
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), link_font)
        rFonts.set(qn('w:hAnsi'), link_font)
        rFonts.set(qn('w:cs'), link_font)
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(link_size_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(link_size_pt * 2)))
        rPr.append(szCs)
        
        run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)
        hyperlink.append(run)
        link_p.append(hyperlink)
        insert_para_xml(link_p)
        
    pb_para2 = OxmlElement('w:p')
    run2 = OxmlElement('w:r')
    br2 = OxmlElement('w:br')
    br2.set(qn('w:type'), 'page')
    run2.append(br2)
    pb_para2.append(run2)
    insert_para_xml(pb_para2)
