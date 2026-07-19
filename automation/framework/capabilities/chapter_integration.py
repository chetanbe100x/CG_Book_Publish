from __future__ import annotations

import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..capabilities.legacy_fonts import is_legacy_font, KrutiDev_to_Unicode, convert_runs_to_unicode, convert_table_runs

def iter_block_members(parent):
    parent_elm = parent.element.body if (hasattr(parent, "element") and hasattr(parent.element, "body")) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

def get_unicode_text(p: Paragraph) -> str:
    parts = []
    for run in p.runs:
        font_name = run.font.name
        run_text = run.text
        if is_legacy_font(font_name) and run_text:
            parts.append(KrutiDev_to_Unicode(run_text))
        else:
            parts.append(run_text or "")
    return "".join(parts).strip()

def detect_chapter_ranges(ref_doc: Document, chapters_config: list[dict]) -> list[dict]:
    unicode_texts = [get_unicode_text(p) for p in ref_doc.paragraphs]
    ranges = []
    
    for cfg in chapters_config:
        start_pat = cfg.get("ref_start_pattern")
        end_pat = cfg.get("ref_end_pattern")
        
        start_idx = -1
        end_idx = -1
        
        # Search for start pattern
        for idx, text in enumerate(unicode_texts):
            if not text:
                continue
            if re.search(start_pat, text, re.IGNORECASE):
                start_idx = idx
                break
                
        # Search for end pattern starting from start_idx
        if start_idx != -1:
            if end_pat == "$END":
                end_idx = len(unicode_texts) - 1
            else:
                for idx in range(start_idx + 1, len(unicode_texts)):
                    text = unicode_texts[idx]
                    if not text:
                        continue
                    if re.search(end_pat, text, re.IGNORECASE):
                        end_idx = idx - 1  # end index is inclusive of the element before the end match
                        break
                if end_idx == -1:
                    end_idx = len(unicode_texts) - 1
                    
        ranges.append({
            "name": cfg.get("name"),
            "start_idx": start_idx,
            "end_idx": end_idx,
            "unified_title": cfg.get("unified_title"),
            "insert_before_pattern": cfg.get("insert_before_pattern"),
            "bilingual_pairing": cfg.get("bilingual_pairing", [])
        })
    return ranges

def detect_insertion_point(target_doc: Document, insert_before_pattern: str) -> int:
    for idx, p in enumerate(target_doc.paragraphs):
        text = p.text.strip()
        # Find chapter matching the pattern
        if re.search(insert_before_pattern, text, re.IGNORECASE):
            return idx
    return 181  # Fallback to the original hardcoded index if not found

def filter_separator_lines(text: str) -> bool:
    text_strip = text.strip()
    if text_strip and all(char in "-_* " for char in text_strip) and len(text_strip) > 5:
        return True
    return False

def convert_paragraph_to_unicode(src_p, dest_doc):
    dest_p = dest_doc.add_paragraph()
    dest_p.alignment = src_p.alignment
    dest_p.paragraph_format.left_indent = src_p.paragraph_format.left_indent
    dest_p.paragraph_format.right_indent = src_p.paragraph_format.right_indent
    dest_p.paragraph_format.space_before = src_p.paragraph_format.space_before
    dest_p.paragraph_format.space_after = src_p.paragraph_format.space_after
    dest_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing
    
    # Process runs using legacy conversion helper
    convert_runs_to_unicode(src_p, dest_p)
    return dest_p

def convert_table_element(src_table):
    import copy
    tbl_clone = copy.deepcopy(src_table._tbl)
    # Translate KrutiDev/Kundli text/fonts to Mangal inside XML elements of table
    convert_table_runs(tbl_clone)
    return tbl_clone
