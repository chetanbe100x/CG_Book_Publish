from __future__ import annotations

import os
import re
import subprocess
import json
import hashlib
from pathlib import Path
from typing import Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..core.models import JobConfig, FrameworkError, sha256_file
from ..capabilities.legacy_fonts import is_legacy_font, KrutiDev_to_Unicode, convert_runs_to_unicode, convert_table_runs
from ..capabilities.chapter_integration import iter_block_members, get_unicode_text, detect_chapter_ranges, detect_insertion_point, filter_separator_lines, convert_paragraph_to_unicode, convert_table_element
from ..capabilities.chapter_structure import relocate_chapter, unify_headings
from ..capabilities.toc_generator import generate_toc

def run_git_checkout(file_path: Path) -> None:
    try:
        subprocess.run(["git", "checkout", "--", str(file_path)], capture_output=True, check=True)
        print(f"[reference] Discarded local changes for {file_path.name} via Git.")
    except Exception as e:
        print(f"[reference] Git checkout warning (non-fatal): {e}")

def run_word_pagination(doc_path: Path, output_path: Path) -> None:
    print(f"[reference] Running Word pagination on {doc_path.name}...")
    try:
        import win32com.client
    except ImportError as e:
        raise FrameworkError("Word pagination requires pywin32 (pip install pywin32)") from e
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(str(doc_path))
        doc.Repaginate()
        
        paragraphs = doc.Paragraphs
        page_offsets = []
        
        for idx in range(1, paragraphs.Count + 1):
            p = paragraphs.Item(idx)
            page_num = p.Range.Information(3)  # wdActiveEndPageNumber
            page_offsets.append(page_num)
            
        doc.Close(False)
        word.Quit()
        
        # Save output document
        import shutil
        shutil.copyfile(doc_path, output_path)
        
        # Write metadata file for pagination
        meta_path = output_path.with_suffix(".docx.metadata.json")
        import json
        meta_data = {"page_offsets": page_offsets}
        meta_path.write_text(json.dumps(meta_data), encoding="utf-8")
        print(f"[reference] Pagination data saved to {meta_path.name}")
    except Exception as e:
        if word:
            try:
                word.Quit()
            except:
                pass
        raise FrameworkError(f"Word pagination failed: {e}")

def copy_table_element_reconstruct(src_table, dest_doc):
    import copy
    tbl_clone = copy.deepcopy(src_table._tbl)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for r in tbl_clone.findall(".//w:r", namespaces=ns):
        text_nodes = r.findall("w:t", namespaces=ns)
        if text_nodes:
            text_val = "".join([t.text for t in text_nodes if t.text])
            is_hindi = any(ord(char) > 127 for char in text_val)
            
            rPr = r.get_or_add_rPr()
            sz = rPr.find("w:sz", namespaces=ns)
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), "23" if is_hindi else "21")
            
            szCs = rPr.find("w:szCs", namespaces=ns)
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rPr.append(szCs)
            szCs.set(qn("w:val"), "23" if is_hindi else "21")
            
            rFonts = rPr.find("w:rFonts", namespaces=ns)
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            if is_hindi:
                rFonts.set(qn("w:hint"), "default")
                rFonts.set(qn("w:ascii"), "Mangal")
                rFonts.set(qn("w:hAnsi"), "Mangal")
                rFonts.set(qn("w:cs"), "Mangal")
            else:
                rFonts.set(qn("w:hint"), "default")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rFonts.set(qn("w:cs"), "Times New Roman")
                
    for tr in tbl_clone.findall(".//w:tr", namespaces=ns):
        trPr = tr.get_or_add_trPr()
        if trPr.find("w:cantSplit", namespaces=ns) is None:
            trPr.append(OxmlElement("w:cantSplit"))
            
    dest_doc.element.body.append(tbl_clone)

def set_run_font_complex(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

def add_mixed_run(paragraph, text, is_bold=False, is_italic=False, size_hindi=12.0, size_latin=11.0):
    tokens = []
    current_token = ""
    current_type = None
    
    for char in text:
        is_hindi_char = ('\u0900' <= char <= '\u097f')
        char_type = "hindi" if is_hindi_char else "latin"
        
        if char in " \t\r\n,.:;?!()₹%-0123456789/\\":
            if current_type is None:
                current_type = char_type
            char_type = current_type
            
        if char_type != current_type:
            if current_token:
                tokens.append((current_token, current_type))
            current_token = char
            current_type = char_type
        else:
            current_token += char
            
    if current_token:
        tokens.append((current_token, current_type))
        
    for token_text, token_type in tokens:
        run = paragraph.add_run(token_text)
        run.bold = is_bold
        run.italic = is_italic
        if token_type == "hindi":
            set_run_font_complex(run, "Mangal")
            run.font.size = Pt(size_hindi)
        else:
            set_run_font_complex(run, "Times New Roman")
            run.font.size = Pt(size_latin)

def create_options_table(doc, options):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = trPr.find(qn("w:cantSplit"))
        if cantSplit is None:
            trPr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.width = Inches(3.0)
            
    cells = [table.cell(0, 0), table.cell(0, 1), table.cell(1, 0), table.cell(1, 1)]
    for idx, text in enumerate(options):
        p = cells[idx].paragraphs[0]
        p.style = "Normal"
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        
        match = re.match(r"^\(?([a-dE-He-hA-Da-d])\)?\s*(.*)", text)
        if match:
            label = match.group(1)
            val = match.group(2)
            run_lbl = p.add_run(f"({label}) ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = RGBColor(0x16, 0x88, 0x73)
            set_run_font_complex(run_lbl, "Times New Roman")
            run_lbl.font.size = Pt(11.0)
            add_mixed_run(p, val, size_hindi=11.5, size_latin=10.5)
        else:
            add_mixed_run(p, text, size_hindi=11.5, size_latin=10.5)

def add_dotted_line(doc, index=0, is_last=False, keep_with_next=False, spacing_points=30.5):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(0)
    if is_last:
        p.paragraph_format.space_after = Pt(20.0)
    else:
        p.paragraph_format.space_after = Pt(max(1.0, spacing_points - 11.0))
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep_with_next
    
    pPr = p._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "4")
    space_val = "1" if index % 2 == 0 else "2"
    bottom.set(qn("w:space"), space_val)
    bottom.set(qn("w:color"), "B7B7B7")

    between = borders.find(qn("w:between"))
    if between is None:
        between = OxmlElement("w:between")
        borders.append(between)
    between.set(qn("w:val"), "dotted")
    between.set(qn("w:sz"), "4")
    between.set(qn("w:space"), space_val)
    between.set(qn("w:color"), "B7B7B7")

def run_reconstruction(paginated_path: Path, output_path: Path, metadata_path: Path) -> None:
    print(f"[reference] Running layout reconstruction on {paginated_path.name}...")
    import json
    
    doc = Document(paginated_path)
    out_doc = Document()
    
    # Setup page margins
    for section in out_doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    line_count = 0
    is_mcq = False
    
    # Iterate through elements (paragraphs & tables) sequentially
    elements = list(iter_block_members(doc))
    
    i = 0
    while i < len(elements):
        member = elements[i]
        
        # If Table, copy element directly (preserves formatting and properties)
        if isinstance(member, Table):
            copy_table_element_reconstruct(member, out_doc)
            i += 1
            continue
            
        p = member
        text = p.text.replace("\u200b", "").replace("\ufeff", "").strip()
        if not text:
            i += 1
            continue
            
        # Check if heading
        is_heading = False
        if re.search(r"^(अध्याय|Chapter|CHAPTER|भाग|Part|Part-)", text):
            is_heading = True
        elif len(text) < 120 and any(kw in text for kw in ["Marks", "numerical", "Objective", "Theory", "Questions", "Question"]):
            is_heading = True
            
        if is_heading:
            # Try to match line count directives
            line_match = re.search(r"प्रत्येक\s+प्रश्न\s*(पर|में)\s*(\d+)\s*लाइन", text)
            if line_match:
                line_count = int(line_match.group(2))
                is_mcq = False
            else:
                line_match = re.search(r"(\d+)\s*(लाइन|line)", text)
                if line_match:
                    line_count = int(line_match.group(1))
                    is_mcq = False
                else:
                    # Check if it specifies MCQ
                    if any(kw in text.casefold() for kw in ["वस्तुनिष्ठ", "mcq", "multiple choice", "वैकल्पिक", "objective"]):
                        line_count = 0
                        is_mcq = True
                    else:
                        # Apply defaults for specific types if line count not explicitly written
                        if any(kw in text.casefold() for kw in ["numerical", "practical", "क्रियात्मक", "न्यूमेरिकल"]):
                            line_count = 10
                        elif any(kw in text.casefold() for kw in ["theory", "descriptive", "वर्णनात्मक", "सिद्धांतिक"]):
                            line_count = 8
                        else:
                            line_count = 0
                        is_mcq = False
                    
            print(f"Header: '{text}' -> Classified: MCQ={is_mcq}, Dotted Lines={line_count}".encode("ascii", errors="backslashreplace").decode("ascii"))
            
            # Write heading
            h_p = out_doc.add_paragraph()
            is_chapter = "अध्याय" in text or "Chapter" in text
            if is_chapter:
                h_p.alignment = 1 # Center
                h_p.paragraph_format.page_break_before = True # Always start new chapter on a new page!
                add_mixed_run(h_p, text, is_bold=True, size_hindi=15.0, size_latin=14.0)
            else:
                h_p.alignment = 0 # Left
                add_mixed_run(h_p, text, is_bold=True, size_hindi=13.0, size_latin=12.0)
                
            i += 1
            continue
            
        # Check if question
        is_q = re.match(r"^(Q\d+\.|प्रश्न\s*:\s*\d*|प्रश्न\s+\d+|Q\d+)", text)
        if is_q:
            q_label = is_q.group(1)
            q_text = text[len(q_label):].strip()
            
            # Gather English translation if immediately following
            eng_text = ""
            if i + 1 < len(elements) and isinstance(elements[i+1], Paragraph):
                next_p = elements[i+1]
                next_text = next_p.text.replace("\u200b", "").replace("\ufeff", "").strip()
                # If next line is not a question and not an option, treat as English translation
                if next_text and not re.match(r"^(Q\d+\.|प्रश्न|[\u200b\s]*\([a-dE-He-hA-Da-d]\))", next_text):
                    eng_text = next_text
                    i += 1 # Consume English paragraph
                    
            # Write Hindi Question
            q_p = out_doc.add_paragraph()
            q_p.paragraph_format.keep_together = True
            q_p.paragraph_format.keep_with_next = True
            
            run_lbl = q_p.add_run(f"{q_label} ")
            run_lbl.bold = True
            run_lbl.font.color.rgb = RGBColor(0x16, 0x88, 0x73) # Accent green
            set_run_font_complex(run_lbl, "Times New Roman")
            run_lbl.font.size = Pt(11.5)
            add_mixed_run(q_p, q_text, size_hindi=11.5, size_latin=11.0)
            
            # Write English Translation
            if eng_text:
                eng_p = out_doc.add_paragraph()
                eng_p.paragraph_format.left_indent = Pt(12) # Indent translation
                eng_p.paragraph_format.keep_together = True
                eng_p.paragraph_format.keep_with_next = True
                add_mixed_run(eng_p, eng_text, size_hindi=11.0, size_latin=10.0)
                for run in eng_p.runs:
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate gray
                
            # If MCQ section or options are immediately following, gather options
            next_cand_text = ""
            if i + 1 < len(elements) and isinstance(elements[i+1], Paragraph):
                next_cand_text = elements[i+1].text.replace("\u200b", "").strip()
                
            if is_mcq or re.match(r"^\([a-dE-He-hA-Da-d]\)", next_cand_text):
                options = []
                opt_count = 0
                while opt_count < 4 and i + 1 < len(elements):
                    cand_el = elements[i+1]
                    if not isinstance(cand_el, Paragraph):
                        break
                    cand_text = cand_el.text.replace("\u200b", "").replace("\ufeff", "").strip()
                    if cand_text and re.match(r"^\([a-dE-He-hA-Da-d]\)", cand_text):
                        options.append(cand_text)
                        opt_count += 1
                        i += 1
                    elif not cand_text:
                        i += 1 # skip empty lines
                    else:
                        break
                        
                if len(options) == 4:
                    create_options_table(out_doc, options)
                else:
                    # Write whatever options we gathered
                    for opt in options:
                        opt_p = out_doc.add_paragraph()
                        opt_p.paragraph_format.keep_together = True
                        opt_p.paragraph_format.keep_with_next = True
                        opt_p.add_run(opt)
                # MCQ: Append exactly 2 dotted lines after options grid with gap
                for line_idx in range(2):
                    is_last_line = (line_idx == 1)
                    keep = not is_last_line
                    add_dotted_line(out_doc, line_idx, is_last=is_last_line, keep_with_next=keep)
            else:
                # Non-MCQ question: append dotted answer lines
                if line_count > 0:
                    for line_idx in range(line_count):
                        is_last_line = (line_idx == line_count - 1)
                        if line_count <= 4:
                            keep = not is_last_line
                        else:
                            keep = (line_idx < 2)
                        add_dotted_line(out_doc, line_idx, is_last=is_last_line, keep_with_next=keep)
                        
            i += 1
            continue
            
        # Default copy for regular body text/instructions
        new_p = out_doc.add_paragraph()
        new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
        new_p.alignment = p.alignment
        add_mixed_run(new_p, text, size_hindi=12.0, size_latin=11.0)
                
        i += 1
        
    out_doc.save(output_path)
    print(f"[reference] Reconstruction successfully written to {output_path.name}")

def run_fake_pagination(source_path: Path, target_path: Path) -> None:
    print(f"[reference] Running fake pagination on {source_path.name}...")
    doc = Document(source_path)
    break_count = 0
    for idx, p in enumerate(doc.paragraphs):
        # Insert page break every 10 paragraphs up to 15
        if idx > 0 and idx % 10 == 0 and break_count < 15:
            run = p.add_run()
            rPr = run._r.get_or_add_rPr()
            lrpb = OxmlElement('w:lastRenderedPageBreak')
            rPr.append(lrpb)
            break_count += 1
            
    print(f"[reference] Inserted {break_count} fake page breaks.")
    doc.save(target_path)
    print(f"[reference] Saved final pagination doc to {target_path.name}")

def validate_prepared_reference(job: JobConfig) -> None:
    """Verify that prepared reference exists and matches current input hashes."""
    prepared_path = job.prepared_source
    if prepared_path is None:
        return
    if not prepared_path.is_file():
        raise FrameworkError(
            f"Prepared reference file is missing at: {prepared_path}. "
            "Please run 'prepare-reference' command first."
        )
    state_file = job.qa_dir / "reference_ingest.state"
    if not state_file.is_file():
        raise FrameworkError(
            "Prepared reference state is missing. Please run 'prepare-reference' command."
        )
    try:
        stored_state = json.loads(state_file.read_text(encoding="utf-8"))
        clean_source_path = job.source.parent / (job.source.name.replace("_reconstructed_paginated", ""))
        current_state = {
            "layout_reference_sha256": sha256_file(job.layout_reference),
            "source_sha256": sha256_file(clean_source_path),
            "chapter_integration_hash": hashlib.sha256(
                json.dumps(job.chapter_integration, sort_keys=True).encode("utf-8")
            ).hexdigest()
        }
        if (stored_state.get("layout_reference_sha256") != current_state["layout_reference_sha256"] or
            stored_state.get("source_sha256") != current_state["source_sha256"] or
            stored_state.get("chapter_integration_hash") != current_state["chapter_integration_hash"]):
            raise FrameworkError(
                "Prepared reference is stale. Please rerun 'prepare-reference' command."
            )
    except Exception as exc:
        if isinstance(exc, FrameworkError):
            raise
        raise FrameworkError(f"Failed to validate prepared reference freshness: {exc}") from exc


def build_ingested_reference(job: JobConfig) -> None:
    # 1. Derive original clean source path
    clean_source_path = job.source.parent / (job.source.name.replace("_reconstructed_paginated", ""))
    
    # 2. Re-read layout reference from config
    layout_ref_path = job.layout_reference
    if not layout_ref_path or not layout_ref_path.is_file():
        raise FrameworkError(f"Layout reference not configured or not found: {layout_ref_path}")
        
    # Check freshness of layout reference to avoid redundant build
    state_file = job.qa_dir / "reference_ingest.state"
    current_state = {
        "layout_reference_sha256": sha256_file(layout_ref_path),
        "source_sha256": sha256_file(clean_source_path),
        "chapter_integration_hash": hashlib.sha256(
            json.dumps(job.chapter_integration, sort_keys=True).encode("utf-8")
        ).hexdigest()
    }
    
    prepared_path = job.prepared_source
    assert prepared_path is not None
    is_fresh = False
    if state_file.is_file() and prepared_path.is_file():
        try:
            stored_state = json.loads(state_file.read_text(encoding="utf-8"))
            if (stored_state.get("layout_reference_sha256") == current_state["layout_reference_sha256"] and
                stored_state.get("source_sha256") == current_state["source_sha256"] and
                stored_state.get("chapter_integration_hash") == current_state["chapter_integration_hash"]):
                is_fresh = True
        except Exception:
            pass
            
    if is_fresh:
        print(f"[reference] Ingested reference is already up-to-date. Skipping build.")
        return
        
    print(f"[reference] Starting reference ingestion pipeline for subject: {job.subject}...")
    
    # 3. Copy clean source to work directory to start fresh (Input/ remains untouched)
    raw_source_path = job.work_dir / clean_source_path.name
    job.work_dir.mkdir(parents=True, exist_ok=True)
    import shutil
    shutil.copyfile(clean_source_path, raw_source_path)
    
    # 4. Integrate Chapters
    ref_doc = Document(layout_ref_path)
    target_doc = Document(raw_source_path)
    
    # Resolve ranges and insertion points via content-based regex matching
    chapters_config = job.chapter_integration.get("chapters", [])
    resolved_chapters = detect_chapter_ranges(ref_doc, chapters_config)
    
    temp_doc = Document()
    q_counter = 0
    
    # Extract blocks and build integrated sequence
    for r_cfg in resolved_chapters:
        start_idx = r_cfg["start_idx"]
        end_idx = r_cfg["end_idx"]
        unified_title = r_cfg["unified_title"]
        insert_before_pat = r_cfg["insert_before_pattern"]
        bilingual_pairing = r_cfg["bilingual_pairing"]
        
        if start_idx == -1 or end_idx == -1:
            print(f"[reference] Warning: Chapter range '{r_cfg['name']}' not detected in layout reference.")
            continue
            
        print(f"[reference] Ingesting chapter '{r_cfg['name']}' (reference index range: {start_idx}-{end_idx})...")
        
        # Iterate and copy reference blocks into temp_doc
        paragraph_idx = 0
        chapter_blocks = []
        for member in iter_block_members(ref_doc):
            if isinstance(member, Paragraph):
                if start_idx <= paragraph_idx <= end_idx:
                    chapter_blocks.append((member, "paragraph", paragraph_idx))
                paragraph_idx += 1
            elif isinstance(member, Table):
                preceding_idx = paragraph_idx - 1
                if start_idx <= preceding_idx <= end_idx:
                    chapter_blocks.append((member, "table", preceding_idx))
                    
        # Apply titles, filters and pairing logic
        # 1. Insert Unified Heading at start of chapter
        dest_p = temp_doc.add_paragraph()
        h_run = dest_p.add_run(unified_title)
        h_run.bold = True
        h_run.font.name = "Mangal"
        
        # Bilingual pairing variables
        hindi_adjustments = []
        english_adjustments = []
        
        for block, b_type, p_idx in chapter_blocks:
            if b_type == "table":
                temp_doc.element.body.append(convert_table_element(block))
            else:
                text = block.text.strip()
                if filter_separator_lines(text):
                    continue
                    
                # Pair matching lists (e.g. revaluation adjustments)
                is_paired = False
                for pair_cfg in bilingual_pairing:
                    p_count = pair_cfg["pair_count"]
                    h_start_idx = pair_cfg.get("hindi_start_index")
                    h_end_idx = pair_cfg.get("hindi_end_index")
                    e_start_idx = pair_cfg.get("english_start_index")
                    e_end_idx = pair_cfg.get("english_end_index")
                    
                    if h_start_idx is not None and h_end_idx is not None and h_start_idx <= p_idx <= h_end_idx:
                        hindi_adjustments.append((block, p_idx))
                        is_paired = True
                        break
                    elif e_start_idx is not None and e_end_idx is not None and e_start_idx <= p_idx <= e_end_idx:
                        english_adjustments.append((block, p_idx))
                        is_paired = True
                        if p_idx == e_end_idx:
                            print("[reference] Pairing bilingual adjustments...")
                            convert_paragraph_to_unicode(hindi_adjustments[0][0], temp_doc)
                            convert_paragraph_to_unicode(english_adjustments[0][0], temp_doc)
                            for k in range(1, p_count + 1):
                                q_counter += 1
                                h_p = hindi_adjustments[k][0]
                                e_p = english_adjustments[k][0]
                                
                                dest_p = temp_doc.add_paragraph()
                                dest_p.alignment = h_p.alignment
                                dest_p.paragraph_format.left_indent = h_p.paragraph_format.left_indent
                                run_lbl = dest_p.add_run(f"Q{q_counter}. ")
                                run_lbl.bold = True
                                convert_runs_to_unicode(h_p, dest_p)
                                
                                dest_p2 = temp_doc.add_paragraph()
                                dest_p2.alignment = e_p.alignment
                                dest_p2.paragraph_format.left_indent = e_p.paragraph_format.left_indent
                                convert_runs_to_unicode(e_p, dest_p2)
                                
                            convert_paragraph_to_unicode(hindi_adjustments[p_count + 1][0], temp_doc)
                            convert_paragraph_to_unicode(english_adjustments[p_count + 1][0], temp_doc)
                        break
                        
                if is_paired:
                    continue
                    
                # Reset question numbering at section headers
                if "vad ds" in text or "Marks" in text or "practical" in text.casefold() or "Theory" in text or "Numerical" in text or "Question" in text:
                    q_counter = 0
                    
                # Check for auto-numbered list item
                pPr = block._p.pPr
                is_list = False
                if pPr is not None:
                    numPr = pPr.find(qn("w:numPr"))
                    is_list = numPr is not None
                    
                if is_list and text and not ("अध्याय" in text or "chapter" in text.lower() or "v/;k;" in text.lower()):
                    q_counter += 1
                    dest_p = temp_doc.add_paragraph()
                    dest_p.alignment = block.alignment
                    dest_p.paragraph_format.left_indent = block.paragraph_format.left_indent
                    
                    run_lbl = dest_p.add_run(f"Q{q_counter}. ")
                    run_lbl.bold = True
                    convert_runs_to_unicode(block, dest_p)
                else:
                    convert_paragraph_to_unicode(block, temp_doc)
                    
    # Locate Chapter 11 sequence correction inside target_doc
    struct_cfg = job.chapter_structure
    chapter_move = struct_cfg.get("logical_order", [])
    # Relocate misplaced chapters (Ch 11 relocates before Part B)
    relocate_chapter(target_doc, "रोकड़.*प्रवाह|Cash Flow", "भाग-ब|Part B")
    
    # Unify target document headings
    font_map = dict(job.pdf_font_map)
    target_font = font_map.get("hindi", "Mangal")
    heading_unifications = struct_cfg.get("heading_unification", [])
    unify_headings(target_doc, heading_unifications, target_font=target_font)
    
    # Insert new chapters into target_doc
    insertion_pat = resolved_chapters[0]["insert_before_pattern"]
    insertion_idx = detect_insertion_point(target_doc, insertion_pat)
    safe_insert_pat = insertion_pat.encode('ascii', errors='backslashreplace').decode('ascii')
    print(f"[reference] Inserting chapters at index {insertion_idx} (before '{safe_insert_pat}')...")
    
    target_body = target_doc.element.body
    insert_before_el = target_doc.paragraphs[insertion_idx]._p
    pos = target_body.index(insert_before_el)
    
    temp_elements = list(iter_block_members(temp_doc))
    for el in temp_elements:
        if isinstance(el, Paragraph):
            el_xml = el._p
        elif isinstance(el, Table):
            el_xml = el._tbl
        target_body.insert(pos, el_xml)
        pos += 1
        
    target_doc.save(raw_source_path)
    print(f"[reference] Ingestion step complete. Saved clean source to {raw_source_path.name}")
    
    # 5. Run Word Pagination
    paginated_path = raw_source_path.parent / (raw_source_path.stem + "_paginated.docx")
    run_word_pagination(raw_source_path, paginated_path)
    
    # 6. Run Reconstruction
    reconstructed_path = raw_source_path.parent / (raw_source_path.stem + "_reconstructed.docx")
    meta_path = paginated_path.with_suffix(".docx.metadata.json")
    run_reconstruction(paginated_path, reconstructed_path, meta_path)
    
    # 7. Run Table of Contents Generation
    toc_config = job.toc
    if toc_config.get("enabled", True):
        reconstructed_doc = Document(reconstructed_path)
        generate_toc(reconstructed_doc, toc_config)
        reconstructed_doc.save(reconstructed_path)
        print(f"[reference] Table of Contents added to {reconstructed_path.name}")
        
    # 8. Run Fake Pagination (generating job.prepared_source)
    run_fake_pagination(reconstructed_path, job.prepared_source)
    
    # Write reference ingest state JSON to skip next runs if unchanged
    state_file.parent.mkdir(parents=True, exist_ok=True)
    state_file.write_text(json.dumps(current_state, indent=2), encoding="utf-8")
    print("[reference] Reference layout ingestion pipeline successfully completed!")
