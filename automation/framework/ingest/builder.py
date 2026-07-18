from __future__ import annotations

import os
from math import ceil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..core.models import FrameworkError, IntegrityError, JobConfig


ACCENT = RGBColor(0x16, 0x88, 0x73)
TEXT = RGBColor(0x18, 0x18, 0x18)


def _set_style_font(style: Any, name: str, size: float, *, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def _ensure_styles(document: Any, job: JobConfig) -> None:
    definitions = (
        ("Book Body", job.font_for("latin"), 10.5, False, 2, 1.05),
        ("Book Hindi", job.font_for("hindi"), 11.5, False, 2, 1.05),
        ("Book Question", job.font_for("hindi"), 11.5, False, 2, 1.05),
        ("Book Section", job.font_for("latin"), 11.5, True, 5, 1.0),
        ("Book Chapter", job.font_for("latin"), 14.0, True, 4, 1.0),
        ("Book Equation", job.font_for("math"), 11.5, False, 3, 1.0),
        ("Book Answer Line", job.font_for("latin"), 9.0, False, 0, 1.0),
    )
    for name, font, size, bold, after, line_spacing in definitions:
        if name in document.styles:
            style = document.styles[name]
        else:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(style, font, size, bold=bold)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.keep_together = True


def _clear_template_body(document: Any) -> None:
    body = document._element.body
    section = body.sectPr
    for child in list(body):
        if child is not section:
            body.remove(child)


def _set_run_font(run: Any, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def _apply_run(run: Any, value: dict[str, Any], job: JobConfig) -> None:
    role = str(value.get("language", "latin"))
    if role not in {"hindi", "latin", "math"}:
        role = "latin"
    _set_run_font(run, job.font_for(role), value.get("size_points"))
    run.bold = bool(value.get("bold", False))
    run.italic = bool(value.get("italic", False))
    run.font.superscript = bool(value.get("superscript", False))
    run.font.subscript = bool(value.get("subscript", False))
    color = value.get("color")
    if isinstance(color, str) and len(color.lstrip("#")) == 6:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if role == "hindi":
        language = OxmlElement("w:lang")
        language.set(qn("w:val"), "hi-IN")
        run._element.get_or_add_rPr().append(language)


def _add_text(paragraph: Any, block: dict[str, Any], job: JobConfig) -> None:
    runs = block.get("runs")
    if isinstance(runs, list):
        for value in runs:
            if not isinstance(value, dict):
                raise FrameworkError("Rich text runs must be objects")
            run = paragraph.add_run(str(value.get("text", "")))
            _apply_run(run, value, job)
        return
    run = paragraph.add_run(str(block.get("text", "")))
    _apply_run(
        run,
        {
            "language": block.get("language", "latin"),
            "bold": block.get("bold", False),
            "italic": block.get("italic", False),
            "size_points": block.get("size_points"),
            "color": block.get("color"),
        },
        job,
    )


def _add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _remove_table_borders(table: Any) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_width(cell: Any, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.first_child_found_in("w:tcW")
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:type"), "dxa")
    # OOXML table widths are expressed in twentieths of a point (twips),
    # while python-docx Length.__int__ returns EMUs. Word honors the OOXML
    # value strictly and collapses cells when EMUs are written as twips.
    cell_width.set(qn("w:w"), str(width.twips))


def _add_dotted_bottom_border(paragraph: Any) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "dot")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")


def _add_answer_lines(document: Any, block: dict[str, Any]) -> list[Any]:
    lines: list[Any] = []
    count = int(block["count"])
    spacing = float(block.get("spacing_points", 16))
    for _ in range(count):
        paragraph = document.add_paragraph(style="Book Answer Line")
        paragraph.paragraph_format.space_after = Pt(max(1.0, spacing - 11.0))
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS
        )
        paragraph.add_run("\t")
        _add_dotted_bottom_border(paragraph)
        lines.append(paragraph)
    return lines


def _add_options(document: Any, block: dict[str, Any], job: JobConfig) -> Any:
    items = block["items"]
    columns = int(block.get("columns", 2))
    columns = max(1, min(columns, len(items)))
    rows = ceil(len(items) / columns)
    table = document.add_table(rows=rows, cols=columns)
    table.autofit = False
    _remove_table_borders(table)
    width = 6.0 / columns
    for index, item in enumerate(items):
        row, column = divmod(index, columns)
        cell = table.cell(row, column)
        _set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles["Book Equation"]
        label = paragraph.add_run(f"({item['label']}) ")
        _set_run_font(label, job.font_for("latin"), 10.5)
        label.bold = True
        text = paragraph.add_run(str(item["text"]))
        _set_run_font(text, job.font_for(str(item.get("language", "math"))), 10.5)
    return table


def _set_table_borders(table: Any) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "808080")


def _add_table(document: Any, block: dict[str, Any], job: JobConfig) -> Any:
    rows = block["rows"]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    _set_table_borders(table)
    widths = block.get("column_widths_inches")
    if not isinstance(widths, list) or len(widths) != len(rows[0]):
        widths = [6.0 / len(rows[0])] * len(rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            _set_cell_width(cell, float(widths[column_index]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["Book Body"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            _set_run_font(run, job.font_for("latin"), 9.5)
            run.bold = row_index == 0
    return table


def _add_block(
    document: Any,
    block: dict[str, Any],
    job: JobConfig,
    manifest_dir: Path,
) -> Any:
    kind = block["kind"]
    if kind == "heading":
        style = "Book Chapter" if int(block.get("level", 1)) == 1 else "Book Section"
        paragraph = document.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_text(paragraph, block, job)
        return paragraph
    if kind == "paragraph":
        role = str(block.get("role", "body"))
        style = {
            "section_heading": "Book Section",
            "hindi": "Book Hindi",
            "equation": "Book Equation",
        }.get(role, "Book Body")
        paragraph = document.add_paragraph(style=style)
        if block.get("alignment") == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_text(paragraph, block, job)
        return paragraph
    if kind == "question":
        paragraph = document.add_paragraph(style="Book Question")
        label = paragraph.add_run(str(block["label"]) + " ")
        _set_run_font(label, job.font_for("latin"), 11.5)
        label.bold = True
        label.font.color.rgb = ACCENT
        hindi = str(block.get("hindi", "")).strip()
        if hindi:
            run = paragraph.add_run(hindi)
            _apply_run(
                run,
                {"language": "hindi", "size_points": 10.0},
                job,
            )
        english = str(block.get("english", "")).strip()
        if english:
            english_paragraph = document.add_paragraph(style="Book Body")
            english_paragraph.paragraph_format.left_indent = Inches(0.18)
            run = english_paragraph.add_run(english)
            _set_run_font(run, job.font_for("latin"), 10.5)
        return paragraph
    if kind == "options":
        return _add_options(document, block, job)
    if kind == "equation":
        paragraph = document.add_paragraph(style="Book Equation")
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if block.get("alignment", "center") == "center"
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        _add_text(paragraph, {**block, "language": "math"}, job)
        return paragraph
    if kind == "figure":
        path = (manifest_dir / str(block["path"])).resolve()
        if not path.is_file():
            raise FrameworkError(f"Figure is missing: {path}")
        paragraph = document.add_paragraph(style="Book Body")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(float(block.get("width_inches", 3.0))))
        return paragraph
    if kind == "table":
        return _add_table(document, block, job)
    if kind == "spacer":
        paragraph = document.add_paragraph(style="Book Body")
        paragraph.paragraph_format.space_after = Pt(float(block["points"]))
        return paragraph
    if kind == "answer_lines":
        return _add_answer_lines(document, block)
    raise FrameworkError(f"Unsupported Book IR block: {kind}")


def build_normalized_docx(
    job: JobConfig,
    manifest: dict[str, Any],
    output: Path,
) -> dict[str, Any]:
    if output.exists():
        raise IntegrityError(f"Refusing to overwrite normalized source: {output}")
    assert job.content_manifest is not None
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(job.template))
    _clear_template_body(document)
    _ensure_styles(document, job)
    document.core_properties.title = str(manifest["book_name"])
    document.core_properties.subject = "Reviewed PDF normalization source"
    manifest_dir = job.content_manifest.parent
    for page_index, page in enumerate(manifest["pages"]):
        first_element = None
        for block in page["blocks"]:
            created = _add_block(document, block, job, manifest_dir)
            if first_element is None:
                first_element = created
        if first_element is None:
            raise FrameworkError(f"Source page {page['source_page']} produced no content")
        if hasattr(first_element, "_p"):
            _add_bookmark(
                first_element,
                f"SourcePdfPage{int(page['source_page']):03d}",
                page_index + 1,
            )
        if page_index < len(manifest["pages"]) - 1:
            paragraph = document.add_paragraph(style="Book Body")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    temporary = output.with_suffix(output.suffix + ".tmp")
    try:
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        if temporary.exists():
            temporary.unlink()
    return {
        "output": str(output),
        "page_count": len(manifest["pages"]),
        "source_pages": [int(page["source_page"]) for page in manifest["pages"]],
    }
