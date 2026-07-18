from __future__ import annotations

import os
from dataclasses import dataclass
from math import ceil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from ..capabilities.images_vml import set_picture_alt_text, validate_effective_dpi
from ..capabilities.tables import set_exact_geometry
from ..core.layout import PrintLayoutProfile, StyleSpec, load_print_layout, mm_to_twips
from ..core.models import FrameworkError, IntegrityError, JobConfig


STYLE_ROLES = {
    "Book Body": "body",
    "Book Hindi": "hindi",
    "Book English": "english",
    "Book Question": "question",
    "Book Unit": "unit",
    "Book Chapter": "chapter",
    "Book Question Type": "question_type",
    "Book Equation": "equation",
    "Book Option": "option",
    "Book Answer Line": "answer_line",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
CANVAS_ROUNDING_TOLERANCE_TWIPS = 10



@dataclass(frozen=True)
class PageContext:
    left_indent_twips: int
    right_indent_twips: int
    top_offset_twips: int


def _set_style_font(style: Any, name: str, spec: StyleSpec) -> None:
    style.font.name = name
    style.font.size = Pt(spec.size_points)
    style.font.bold = spec.bold
    style.font.italic = spec.italic
    style.font.color.rgb = RGBColor.from_string(spec.color)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def _ensure_styles(
    document: Any,
    job: JobConfig,
    layout: PrintLayoutProfile,
) -> None:
    for style_name, role in STYLE_ROLES.items():
        spec = layout.style(role)
        if style_name in document.styles:
            style = document.styles[style_name]
        else:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(style, job.font_for(spec.font_role), spec)
        paragraph = style.paragraph_format
        paragraph.alignment = ALIGNMENTS[spec.alignment]
        paragraph.space_before = Pt(spec.space_before_points)
        paragraph.space_after = Pt(spec.space_after_points)
        paragraph.line_spacing = spec.line_spacing
        paragraph.left_indent = Twips(spec.left_indent_twips)
        paragraph.right_indent = Twips(0)
        # Page flow is decided from semantic block transitions, not globally.
        paragraph.keep_together = None
        paragraph.keep_with_next = None


def _clear_template_body(document: Any) -> None:
    body = document._element.body
    section = body.sectPr
    for child in list(body):
        if child is not section:
            body.remove(child)


def _configure_document_settings(document: Any, job: JobConfig | None = None) -> None:
    settings = document.settings._element
    for element in list(settings.findall(qn("w:updateFields"))):
        settings.remove(element)
    compression = settings.find(qn("w:doNotAutoCompressPictures"))
    if compression is None:
        compression = OxmlElement("w:doNotAutoCompressPictures")
        settings.append(compression)
        successor = next(
            (
                settings.find(qn(f"w:{name}"))
                for name in ("shapeDefaults", "decimalSymbol", "listSeparator")
                if settings.find(qn(f"w:{name}")) is not None
            ),
            None,
        )
        if successor is not None:
            successor.addprevious(compression)
    compression.set(qn("w:val"), "true")

    if job is not None and job.source_style_policy == "content_only":
        mirror = settings.find(qn("w:mirrorMargins"))
        if mirror is None:
            mirror = OxmlElement("w:mirrorMargins")
            settings.append(mirror)


def _set_run_font(run: Any, name: str, size_points: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_points)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def _apply_run(
    run: Any,
    value: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    base: StyleSpec,
) -> None:
    role = str(value.get("language", base.font_role))
    if role not in {"hindi", "latin", "math"}:
        role = base.font_role
    default_size = layout.inline_math_points if role == "math" else base.size_points
    size = float(value.get("size_points", default_size))
    _set_run_font(run, job.font_for(role), size)
    run.bold = bool(value.get("bold", base.bold))
    run.italic = bool(value.get("italic", base.italic))
    run.font.superscript = bool(value.get("superscript", False))
    run.font.subscript = bool(value.get("subscript", False))
    color = str(value.get("color", base.color)).lstrip("#").upper()
    if len(color) != 6:
        raise FrameworkError(f"Invalid run RGB colour: {color}")
    run.font.color.rgb = RGBColor.from_string(color)
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "hi-IN" if role == "hindi" else "en-IN")
    run._element.get_or_add_rPr().append(language)


def _add_text(
    paragraph: Any,
    block: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    role: str,
    *,
    manifest_dir: Path,
    image_checks: list[dict[str, Any]],
) -> None:
    base = layout.style(role)
    runs = block.get("runs")
    if isinstance(runs, list):
        for value in runs:
            if not isinstance(value, dict):
                raise FrameworkError("Rich text runs must be objects")
            formula = value.get("formula_image")
            if formula is not None:
                if not isinstance(formula, dict):
                    raise FrameworkError("Rich-text formula_image must be an object")
                _add_picture(
                    paragraph,
                    formula,
                    manifest_dir=manifest_dir,
                    layout=layout,
                    kind="formula",
                    image_checks=image_checks,
                )
                continue
            run = paragraph.add_run(str(value.get("text", "")))
            _apply_run(run, value, job, layout, base)
        return
    run = paragraph.add_run(str(block.get("text", "")))
    _apply_run(
        run,
        {
            "language": block.get("language", base.font_role),
            "bold": block.get("bold", base.bold),
            "italic": block.get("italic", base.italic),
            "size_points": block.get("size_points", base.size_points),
            "color": block.get("color", base.color),
        },
        job,
        layout,
        base,
    )


def _add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    properties = paragraph._p.get_or_add_pPr()
    properties_index = paragraph._p.index(properties)
    paragraph._p.insert(properties_index + 1, start)
    paragraph._p.append(end)


def _page_context(
    document: Any,
    layout: PrintLayoutProfile,
    page_index: int,
    first_content_side: str | None = None,
    job: JobConfig | None = None,
) -> PageContext:
    section = document.sections[-1]
    page_width = int(section.page_width.twips)
    page_height = int(section.page_height.twips)
    expected_width = mm_to_twips(layout.page.width_mm)
    expected_height = mm_to_twips(layout.page.height_mm)
    tolerance = mm_to_twips(0.5)
    if abs(page_width - expected_width) > tolerance or abs(page_height - expected_height) > tolerance:
        raise FrameworkError(
            "Template page geometry does not match the print profile: "
            f"{page_width}x{page_height} twips"
        )
    
    if job is not None and job.source_style_policy == "content_only":
        top = max(0, mm_to_twips(layout.page.top_content_mm) - int(section.top_margin.twips))
        return PageContext(left_indent_twips=0, right_indent_twips=0, top_offset_twips=top)

    first_side = first_content_side or layout.page.first_content_side
    if first_side not in {"recto", "verso"}:
        raise FrameworkError("first_content_side must be recto or verso")
    first_is_recto = first_side == "recto"
    is_recto = (page_index % 2 == 0) == first_is_recto
    left_mm = layout.page.inside_margin_mm if is_recto else layout.page.outside_margin_mm
    right_mm = layout.page.outside_margin_mm if is_recto else layout.page.inside_margin_mm
    left = mm_to_twips(left_mm) - int(section.left_margin.twips)
    right = mm_to_twips(right_mm) - int(section.right_margin.twips)
    top = max(0, mm_to_twips(layout.page.top_content_mm) - int(section.top_margin.twips))
    return PageContext(left_indent_twips=left, right_indent_twips=right, top_offset_twips=top)


def _paragraph_role(paragraph: Any) -> str:
    name = getattr(paragraph.style, "name", "Book Body")
    return STYLE_ROLES.get(name, "body")


def _apply_paragraph_geometry(
    paragraph: Any,
    context: PageContext,
    layout: PrintLayoutProfile,
) -> None:
    spec = layout.style(_paragraph_role(paragraph))
    paragraph.paragraph_format.left_indent = Twips(
        context.left_indent_twips + spec.left_indent_twips
    )
    paragraph.paragraph_format.right_indent = Twips(context.right_indent_twips)


def _apply_canvas_width(
    paragraph: Any,
    block: dict[str, Any],
    context: PageContext,
    layout: PrintLayoutProfile,
) -> None:
    width_points = block.get("width_points")
    if width_points is None:
        return
    width_twips = round(float(width_points) * 20)
    live_width = mm_to_twips(layout.page.live_body_width_mm)
    if (
        width_twips <= 0
        or width_twips > live_width + CANVAS_ROUNDING_TOLERANCE_TWIPS
    ):
        raise FrameworkError("answer_canvas.width_points exceeds the live body width")
    width_twips = min(width_twips, live_width)
    remaining = live_width - width_twips
    if block.get("alignment", "center") == "center":
        left_extra = remaining // 2
    else:
        left_extra = 0
    paragraph.paragraph_format.left_indent = Twips(context.left_indent_twips + left_extra)
    paragraph.paragraph_format.right_indent = Twips(
        context.right_indent_twips + remaining - left_extra
    )


def _apply_top_offset(
    paragraph: Any,
    context: PageContext,
    layout: PrintLayoutProfile,
) -> None:
    spec = layout.style(_paragraph_role(paragraph))
    paragraph.paragraph_format.space_before = Pt(
        spec.space_before_points + context.top_offset_twips / 20.0
    )


def _remove_table_borders(table: Any) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_table_borders(table: Any) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "808080")


def _add_paragraph_border(
    paragraph: Any,
    *,
    value: str,
    color: str,
    size: int,
    space: int,
    edges: tuple[str, ...] = ("bottom",),
) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    for edge in edges:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), str(space))
        element.set(qn("w:color"), color)


def _answer_line_count(block: dict[str, Any], layout: PrintLayoutProfile) -> int:
    if "count" in block:
        return int(block["count"])
    marks = int(block.get("marks", 0))
    if marks in layout.answer_space.line_counts_by_marks:
        return layout.answer_space.line_counts_by_marks[marks]
    if marks > 0:
        return max(1, marks * 2)
    raise FrameworkError("answer_lines requires count or marks")


def _answer_line_pitch(block: dict[str, Any], layout: PrintLayoutProfile) -> float:
    if "spacing_points" in block:
        return float(block["spacing_points"])
    if block.get("spacing_mode") == "generous":
        return layout.answer_space.generous_pitch_points
    return layout.answer_space.default_pitch_points


def _add_answer_lines(
    document: Any,
    block: dict[str, Any],
    layout: PrintLayoutProfile,
) -> list[Any]:
    count = _answer_line_count(block, layout)
    pitch = _answer_line_pitch(block, layout)
    if count < 1 or count > 50 or pitch <= 0:
        raise FrameworkError("Invalid answer-line count or pitch")
    lines: list[Any] = []
    for _ in range(count):
        paragraph = document.add_paragraph(style="Book Answer Line")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(pitch)
        paragraph.add_run("\u00A0")
        _add_paragraph_border(
            paragraph,
            value="dotted",
            color=layout.answer_space.border_color,
            size=layout.answer_space.border_size_eighth_points,
            space=layout.answer_space.border_space_points,
            edges=("bottom", "between"),
        )
        lines.append(paragraph)
    return lines


def _formula_visual_units(formula: dict[str, Any]) -> float:
    width_value = formula.get("width_inches")
    if width_value is None:
        bbox = formula.get("crop_bbox_points")
        if isinstance(bbox, list) and len(bbox) == 4:
            try:
                width_value = (float(bbox[2]) - float(bbox[0])) / 72.0
            except (TypeError, ValueError):
                width_value = None
    try:
        width_inches = float(width_value) if width_value is not None else 3.0
    except (TypeError, ValueError):
        width_inches = 3.0
    if width_inches <= 0:
        width_inches = 3.0
    # At 11-14 pt, one printed inch is roughly fourteen average glyphs.
    # This lets compact fractions remain in the source's 2x2 grid while a
    # genuinely wide expression still triggers the one-column safety layout.
    return width_inches * 14.0


def _visual_units(item: dict[str, Any]) -> float:
    formula = item.get("formula_image")
    if isinstance(formula, dict):
        return _formula_visual_units(formula)
    runs = item.get("runs")
    if not isinstance(runs, list):
        text = str(item.get("text", ""))
        return float(len(text) + text.count("\n") * 30)
    units = 0.0
    for run in runs:
        if not isinstance(run, dict):
            continue
        nested = run.get("formula_image")
        if isinstance(nested, dict):
            units += _formula_visual_units(nested)
            continue
        text = str(run.get("text", ""))

        units += len(text) + text.count("\n") * 30
    return units

def _option_columns(block: dict[str, Any], layout: PrintLayoutProfile) -> int:
    items = block["items"]
    requested = int(block.get("columns", 2))
    label_layout = str(block.get("label_layout", "")).casefold()
    if label_layout in {"one_column", "1_column", "stacked"}:
        requested = 1
    elif label_layout in {"two_by_two", "2x2"}:
        requested = 2
    elif label_layout in {"four_inline", "4_inline"}:
        requested = 4
    units = [_visual_units(item) for item in items]
    if any(value > layout.options.max_two_column_visual_units for value in units):
        return 1
    if requested == 4 and len(items) == 4:
        if any(value > layout.options.max_four_inline_visual_units for value in units):
            return 2
        return 4
    if requested == 2 and len(items) >= 2:
        return 2
    return 1


def _image_path(manifest_dir: Path, value: str) -> Path:
    path = (manifest_dir / value).resolve()
    try:
        path.relative_to(manifest_dir.resolve())
    except ValueError as exc:
        raise FrameworkError(f"Image path escapes the manifest directory: {path}") from exc
    if not path.is_file():
        raise FrameworkError(f"Image is missing: {path}")
    return path


def _image_alt_text(block: dict[str, Any], kind: str) -> str:
    explicit = str(block.get("alt_text", "")).strip()
    if explicit:
        return explicit
    source_ids = block.get("source_ids")
    source = str(source_ids[0]) if isinstance(source_ids, list) and source_ids else "source"
    prefix = "Mathematical expression" if kind == "formula" else "Technical figure"
    return f"{prefix} reproduced from {source}"


def _add_picture(
    paragraph: Any,
    block: dict[str, Any],
    *,
    manifest_dir: Path,
    layout: PrintLayoutProfile,
    kind: str,
    image_checks: list[dict[str, Any]],
) -> None:
    path = _image_path(manifest_dir, str(block["path"]))
    width = float(block.get("width_inches", 3.0))
    height_value = block.get("height_inches")
    height = float(height_value) if height_value is not None else None
    if kind == "formula":
        minimum_dpi = layout.images.formula_min_dpi
    elif block.get("image_role") == "decorative":
        minimum_dpi = layout.images.decorative_min_dpi
    else:
        minimum_dpi = layout.images.technical_min_dpi
    dpi = validate_effective_dpi(
        path,
        width_inches=width,
        height_inches=height,
        minimum_dpi=minimum_dpi,
    )
    run = paragraph.add_run()
    kwargs: dict[str, Any] = {"width": Inches(width)}
    if height is not None:
        kwargs["height"] = Inches(height)
    run.add_picture(str(path), **kwargs)
    alt_text = _image_alt_text(block, kind)
    set_picture_alt_text(run, alt_text)
    image_checks.append(
        {
            "kind": kind,
            "path": str(path),
            "effective_dpi": round(dpi, 2),
            "minimum_dpi": minimum_dpi,
            "alt_text": alt_text,
        }
    )


def _add_options(
    document: Any,
    block: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    context: PageContext,
    manifest_dir: Path,
    image_checks: list[dict[str, Any]],
) -> Any:
    items = block["items"]
    columns = _option_columns(block, layout)
    rows = ceil(len(items) / columns)
    table = document.add_table(rows=rows, cols=columns)
    _remove_table_borders(table)
    if columns == 4:
        widths = layout.options.four_inline_widths_twips
    elif columns == 2:
        widths = layout.options.two_column_widths_twips
    else:
        widths = (layout.options.table_width_twips,)
    
    if job.source_style_policy == "content_only":
        total_width = mm_to_twips(layout.page.live_body_width_mm)
        col2_width = total_width - 540
        sum_widths = sum(widths)
        if sum_widths > col2_width:
            scale = col2_width / sum_widths
            widths = tuple(int(w * scale) for w in widths)
        indent = 0
    else:
        indent = context.left_indent_twips + layout.options.table_indent_twips

    set_exact_geometry(
        table,
        widths,
        indent_twips=indent,
        cell_margins_twips=(
            layout.options.cell_margin_top_twips,
            layout.options.cell_margin_right_twips,
            layout.options.cell_margin_bottom_twips,
            layout.options.cell_margin_left_twips,
        ),
    )
    label_spec = layout.style("option_label")
    option_spec = layout.style("option")
    for index, item in enumerate(items):
        row_index, column_index = divmod(index, columns)
        cell = table.cell(row_index, column_index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.style = "Book Option"
        paragraph.paragraph_format.space_after = Pt(layout.options.row_space_after_points)
        label = paragraph.add_run(f"({item['label']}) ")
        _apply_run(label, {}, job, layout, label_spec)
        formula = item.get("formula_image")
        if formula:
            if isinstance(formula, str):
                formula = {"path": formula}
            if not isinstance(formula, dict):
                raise FrameworkError("Option formula_image must be a path or object")
            formula_block = {**block, **formula}
            _add_picture(
                paragraph,
                formula_block,
                manifest_dir=manifest_dir,
                layout=layout,
                kind="formula",
                image_checks=image_checks,
            )
        else:
            _add_text(
                paragraph,
                item,
                job,
                layout,
                "option",
                manifest_dir=manifest_dir,
                image_checks=image_checks,
            )
        # Direct paragraph geometry is intentionally absent inside table cells.
    return table


def _add_table(
    document: Any,
    block: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    context: PageContext,
) -> Any:
    rows = block["rows"]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    _set_table_borders(table)
    widths_value = block.get("column_widths_twips")
    if isinstance(widths_value, list) and len(widths_value) == len(rows[0]):
        widths = tuple(int(value) for value in widths_value)
    else:
        inches_value = block.get("column_widths_inches")
        if isinstance(inches_value, list) and len(inches_value) == len(rows[0]):
            widths = tuple(round(float(value) * 1440) for value in inches_value)
        else:
            total = mm_to_twips(layout.page.live_body_width_mm)
            base, remainder = divmod(total, len(rows[0]))
            widths = tuple(base + (1 if index < remainder else 0) for index in range(len(rows[0])))
    set_exact_geometry(
        table,
        widths,
        indent_twips=context.left_indent_twips + int(block.get("table_indent_twips", 0)),
        cell_margins_twips=(40, 80, 40, 80),
    )
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Book Body"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            _apply_run(
                run,
                {"bold": row_index == 0},
                job,
                layout,
                layout.style("body"),
            )
    return table


def _add_question(
    document: Any,
    block: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    manifest_dir: Path,
    image_checks: list[dict[str, Any]],
) -> list[Any]:
    question = document.add_paragraph(style="Book Question")
    label = question.add_run(str(block["label"]) + " ")
    _apply_run(label, {}, job, layout, layout.style("question_label"))
    created = [question]
    hindi_runs = block.get("hindi_runs")
    hindi = str(block.get("hindi", "")).strip()
    statement = question
    if False and block.get("label_layout", "inline") == "standalone" and (
        (isinstance(hindi_runs, list) and hindi_runs) or hindi
    ):
        statement = document.add_paragraph(style="Book Question")
        question.paragraph_format.keep_with_next = True
        created.append(statement)
    if isinstance(hindi_runs, list) and hindi_runs:
        _add_text(
            statement,
            {"runs": hindi_runs},
            job,
            layout,
            "question",
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )
    elif hindi:
        _add_text(
            statement,
            {"text": hindi, "language": "hindi"},
            job,
            layout,
            "question",
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )

    english_runs = block.get("english_runs")
    english = str(block.get("english", "")).strip()
    if (isinstance(english_runs, list) and english_runs) or english:
        translation = document.add_paragraph(style="Book English")
        payload = (
            {"runs": english_runs}
            if isinstance(english_runs, list) and english_runs
            else {"text": english, "language": "latin"}
        )
        _add_text(
            translation,
            payload,
            job,
            layout,
            "english",
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )
        created[-1].paragraph_format.keep_with_next = True
        created.append(translation)
    return created


def _heading_style(block: dict[str, Any]) -> tuple[str, str]:
    role = str(block.get("role", "")).casefold()
    if role in {"unit", "unit_heading"}:
        return "Book Unit", "unit"
    if role in {"question_type", "question_type_heading", "section_heading"}:
        return "Book Question Type", "question_type"
    if int(block.get("level", 1)) == 1:
        return "Book Chapter", "chapter"
    return "Book Question Type", "question_type"


def _add_answer_canvas(
    document: Any,
    block: dict[str, Any],
    layout: PrintLayoutProfile,
) -> list[Any]:
    if block.get("mode") == "ruled":
        return _add_answer_lines(document, block, layout)
    height = float(block.get("height_points", 180.0))
    if height <= 0:
        raise FrameworkError("answer_canvas.height_points must be positive")
    paragraph = document.add_paragraph(style="Book Answer Line")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(height)
    paragraph.add_run("\u00A0")
    canvas_type = str(block.get("canvas_type", "blank"))
    border_default = canvas_type in {"graph", "venn", "coordinate_grid", "diagram"}
    if bool(block.get("border", border_default)):
        _add_paragraph_border(
            paragraph,
            value="single",
            color=str(block.get("border_color", "A6A6A6")).lstrip("#"),
            size=4,
            space=1,
            edges=("top", "left", "bottom", "right"),
        )
    return [paragraph]


def _add_block(
    document: Any,
    block: dict[str, Any],
    job: JobConfig,
    layout: PrintLayoutProfile,
    context: PageContext,
    manifest_dir: Path,
    image_checks: list[dict[str, Any]],
) -> list[Any]:
    kind = block["kind"]
    if kind == "heading":
        style, role = _heading_style(block)
        paragraph = document.add_paragraph(style=style)
        _add_text(
            paragraph,
            block,
            job,
            layout,
            role,
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )
        return [paragraph]
    if kind == "paragraph":
        role = str(block.get("role", "body"))
        style_role = {
            "section_heading": ("Book Question Type", "question_type"),
            "question_type_heading": ("Book Question Type", "question_type"),
            "hindi": ("Book Hindi", "hindi"),
            "english": ("Book English", "english"),
            "equation": ("Book Equation", "equation"),
        }.get(role, ("Book Body", "body"))
        paragraph = document.add_paragraph(style=style_role[0])
        if block.get("alignment") in ALIGNMENTS:
            paragraph.alignment = ALIGNMENTS[str(block["alignment"])]
        _add_text(
            paragraph,
            block,
            job,
            layout,
            style_role[1],
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )
        return [paragraph]
    if kind == "question":
        return _add_question(document, block, job, layout, manifest_dir, image_checks)
    if kind == "options":
        return [
            _add_options(
                document,
                block,
                job,
                layout,
                context,
                manifest_dir,
                image_checks,
            )
        ]
    if kind == "equation":
        paragraph = document.add_paragraph(style="Book Equation")
        if block.get("alignment") in ALIGNMENTS:
            paragraph.alignment = ALIGNMENTS[str(block["alignment"])]
        _add_text(
            paragraph,
            {**block, "language": "math"},
            job,
            layout,
            "equation",
            manifest_dir=manifest_dir,
            image_checks=image_checks,
        )
        return [paragraph]
    if kind in {"figure", "formula_image"}:
        role = "equation" if kind == "formula_image" else "body"
        style = "Book Equation" if kind == "formula_image" else "Book Body"
        paragraph = document.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_picture(
            paragraph,
            block,
            manifest_dir=manifest_dir,
            layout=layout,
            kind="formula" if kind == "formula_image" else "technical",
            image_checks=image_checks,
        )
        return [paragraph]
    if kind == "table":
        return [_add_table(document, block, job, layout, context)]
    if kind == "spacer":
        paragraph = document.add_paragraph(style="Book Body")
        paragraph.paragraph_format.space_after = Pt(float(block["points"]))
        return [paragraph]
    if kind == "answer_lines":
        return _add_answer_lines(document, block, layout)
    if kind == "answer_canvas":
        return _add_answer_canvas(document, block, layout)
    raise FrameworkError(f"Unsupported Book IR block: {kind}")


def _first_paragraph(elements: list[Any]) -> Any | None:
    for element in elements:
        if hasattr(element, "_p"):
            return element
        if hasattr(element, "rows") and element.rows and element.rows[0].cells:
            return element.rows[0].cells[0].paragraphs[0]
    return None


def _last_paragraph(elements: list[Any]) -> Any | None:
    for element in reversed(elements):
        if hasattr(element, "_p"):
            return element
    return None


def _keep_transition(previous: dict[str, Any], current: dict[str, Any]) -> bool:
    previous_kind = previous.get("kind")
    current_kind = current.get("kind")
    if previous_kind == "heading":
        return True
    if previous_kind == "paragraph" and previous.get("role") in {
        "section_heading",
        "question_type_heading",
    }:
        return True
    if previous_kind == "question" and current_kind in {
        "paragraph",
        "options",
        "equation",
        "formula_image",
        "figure",
        "answer_lines",
        "answer_canvas",
    }:
        current_question = current.get("for_question")
        return current_question in {None, previous.get("id"), previous.get("question_id")}
    return False


def _delete_paragraph(paragraph: Any) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _group_page_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any] | list[dict[str, Any]]]:
    groups: list[Any] = []
    current_q_group: list[dict[str, Any]] | None = None
    for block in blocks:
        kind = block.get("kind")
        if kind == "question":
            current_q_group = [block]
            groups.append(current_q_group)
        elif current_q_group is not None and kind in {
            "options",
            "equation",
            "figure",
            "answer_lines",
            "answer_canvas",
        }:
            current_question = block.get("for_question")
            q_id = current_q_group[0].get("question_id") or current_q_group[0].get("id")
            if current_question in {None, q_id}:
                current_q_group.append(block)
            else:
                current_q_group = None
                groups.append(block)
        else:
            current_q_group = None
            groups.append(block)
    return groups


def build_normalized_docx(
    job: JobConfig,
    manifest: dict[str, Any],
    output: Path,
) -> dict[str, Any]:
    if output.exists():
        raise IntegrityError(f"Refusing to overwrite normalized source: {output}")
    assert job.content_manifest is not None
    output.parent.mkdir(parents=True, exist_ok=True)
    layout = load_print_layout(job.subject_profile)
    document = Document(str(job.template))
    _clear_template_body(document)
    _ensure_styles(document, job, layout)
    _configure_document_settings(document, job)
    if job.source_style_policy == "content_only":
        for section in document.sections:
            section.left_margin = Inches(layout.page.inside_margin_mm / 25.4)
            section.right_margin = Inches(layout.page.outside_margin_mm / 25.4)
            section.top_margin = Inches(layout.page.top_content_mm / 25.4)
            section.bottom_margin = Inches((layout.page.height_mm - layout.page.bottom_content_mm) / 25.4)
    document.core_properties.title = str(manifest["book_name"])
    document.core_properties.subject = "Reviewed PDF normalization source"
    manifest_dir = job.content_manifest.parent
    image_checks: list[dict[str, Any]] = []

    for page_index, page in enumerate(manifest["pages"]):
        context = _page_context(
            document,
            layout,
            page_index,
            getattr(job, "first_content_side", None),
            job=job,
        )
        page_elements: list[Any] = []
        if job.source_style_policy == "content_only":
            groups = _group_page_blocks(page["blocks"])
            for group in groups:
                if isinstance(group, list):
                    q_block = group[0]
                    table = document.add_table(rows=1, cols=2)
                    _remove_table_borders(table)
                    total_width = mm_to_twips(layout.page.live_body_width_mm)
                    col1_width = 540
                    col2_width = total_width - col1_width
                    set_exact_geometry(
                        table,
                        (col1_width, col2_width),
                        indent_twips=context.left_indent_twips,
                        cell_margins_twips=(0, 0, 0, 0),
                    )
                    
                    cell_1 = table.cell(0, 0)
                    p_label = cell_1.paragraphs[0]
                    p_label.style = "Book Question"
                    p_label.paragraph_format.space_after = Pt(layout.style("question").space_after_points)
                    run_label = p_label.add_run(str(q_block.get("label", "")))
                    _apply_run(run_label, {}, job, layout, layout.style("question_label"))
                    
                    cell_2 = table.cell(0, 1)
                    first_p = cell_2.paragraphs[0]
                    
                    cell_context = PageContext(
                        left_indent_twips=0,
                        right_indent_twips=0,
                        top_offset_twips=context.top_offset_twips,
                    )
                    
                    previous_block = None
                    previous_elements = []
                    question_elements = []
                    for idx, block in enumerate(group):
                        if idx == 0:
                            block_to_render = {**block, "label": ""}
                        else:
                            block_to_render = block
                            
                        if previous_block is not None and _keep_transition(previous_block, block_to_render):
                            paragraph = _last_paragraph(previous_elements)
                            if paragraph is not None:
                                paragraph.paragraph_format.keep_with_next = True
                                
                        created = _add_block(
                            cell_2,
                            block_to_render,
                            job,
                            layout,
                            cell_context,
                            manifest_dir,
                            image_checks,
                        )
                        for element in created:
                            if hasattr(element, "_p"):
                                _apply_paragraph_geometry(element, cell_context, layout)
                                if block_to_render.get("kind") == "answer_canvas":
                                    _apply_canvas_width(element, block_to_render, cell_context, layout)
                        question_elements.extend(created)
                        previous_block = block_to_render
                        previous_elements = created
                    
                    if len(cell_2.paragraphs) > 1 and not first_p.text and not first_p.runs:
                        _delete_paragraph(first_p)
                    
                    page_elements.append(table)
                else:
                    created = _add_block(
                        document,
                        group,
                        job,
                        layout,
                        context,
                        manifest_dir,
                        image_checks,
                    )
                    for element in created:
                        if hasattr(element, "_p"):
                            _apply_paragraph_geometry(element, context, layout)
                    page_elements.extend(created)
        else:
            previous_block = None
            previous_elements = []
            for block in page["blocks"]:
                if previous_block is not None and _keep_transition(previous_block, block):
                    paragraph = _last_paragraph(previous_elements)
                    if paragraph is not None:
                        paragraph.paragraph_format.keep_with_next = True
                created = _add_block(
                    document,
                    block,
                    job,
                    layout,
                    context,
                    manifest_dir,
                    image_checks,
                )
                for element in created:
                    if hasattr(element, "_p"):
                        _apply_paragraph_geometry(element, context, layout)
                        if block.get("kind") == "answer_canvas":
                            _apply_canvas_width(element, block, context, layout)
                page_elements.extend(created)
                previous_block = block
                previous_elements = created
        
        anchor = _first_paragraph(page_elements)
        if anchor is None:
            raise FrameworkError(f"Source page {page['source_page']} produced no content")
        top_paragraph = next((item for item in page_elements if hasattr(item, "_p")), None)
        if top_paragraph is not None:
            _apply_top_offset(top_paragraph, context, layout)
        _add_bookmark(
            anchor,
            f"SourcePdfPage{int(page['source_page']):03d}",
            page_index + 1,
        )
        if job.pagination_policy != "sample_flow" and page_index < len(manifest["pages"]) - 1:
            paragraph = document.add_paragraph(style="Book Body")
            _apply_paragraph_geometry(paragraph, context, layout)
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    temporary = output.with_suffix(output.suffix + ".tmp")
    try:
        document.save(temporary)
        os.replace(temporary, output)
    finally:
        if temporary.exists():
            temporary.unlink()
    return {
        "output": str(output),
        "page_count": len(manifest["pages"]),
        "source_pages": [int(page["source_page"]) for page in manifest["pages"]],
        "layout_profile": layout.name,
        "image_checks": image_checks,
    }
